from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Cm


DOC_PATH = Path(r"d:\桌面\课程\节能减排大赛\网站功能文档-用户端与管理员端.docx")
QR_PATH = Path(r"d:\桌面\课程\节能减排大赛\assets\images\site-qrcode-poster.png")


def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def paragraph_has_drawing(paragraph):
    return bool(paragraph._element.xpath(".//w:drawing"))


def main():
    doc = Document(str(DOC_PATH))

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "网站地址：https://134139.xyz/":
            next_p = paragraph._element.getnext()
            if next_p is not None:
                sibling = Paragraph(next_p, paragraph._parent)
                if paragraph_has_drawing(sibling):
                    remove_paragraph(sibling)

            image_paragraph = insert_paragraph_after(paragraph)
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_paragraph.add_run().add_picture(str(QR_PATH), width=Cm(6.2))
            break

    try:
        doc.save(str(DOC_PATH))
    except PermissionError:
        fallback = DOC_PATH.with_name(DOC_PATH.stem + "-带二维码" + DOC_PATH.suffix)
        doc.save(str(fallback))


if __name__ == "__main__":
    main()
