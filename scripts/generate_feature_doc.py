from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = Path(r"d:\桌面\课程\节能减排大赛\网站功能文档-用户端与管理员端.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_border(paragraph, color="9DB8A7", size="8", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right"):
        element = p_bdr.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            p_bdr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), space)
        element.set(qn("w:color"), color)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=12)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=12)


def add_shot_box(doc, title, content, placement):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.35
    set_paragraph_border(p)

    r1 = p.add_run("建议截图：")
    set_run_font(r1, size=11, bold=True, color="0F3D2B")
    r2 = p.add_run(title + "\n")
    set_run_font(r2, size=11)

    r3 = p.add_run("建议截取内容：")
    set_run_font(r3, size=11, bold=True, color="0F3D2B")
    r4 = p.add_run(content + "\n")
    set_run_font(r4, size=11)

    r5 = p.add_run("放置位置：")
    set_run_font(r5, size=11, bold=True, color="0F3D2B")
    r6 = p.add_run(placement)
    set_run_font(r6, size=11)


def add_tip_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.35
    set_paragraph_border(p, color="C1CFC6")
    run1 = p.add_run("文档写作提示：")
    set_run_font(run1, size=11, bold=True, color="0F3D2B")
    run2 = p.add_run(text)
    set_run_font(run2, size=11)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True, color="0F3D2B")
    elif level == 2:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, size=13.5, bold=True, color="0F3D2B")
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, color="0F3D2B")
    return p


def apply_base_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(12)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(12)


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("绿循校园网站功能文档")
    set_run_font(run, size=22, bold=True, color="0F3D2B")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("用户端与管理员端合并说明版")
    set_run_font(run, size=13, color="4B6358")

    doc.add_paragraph()
    add_body_paragraph(doc, "项目名称：绿循校园")
    add_body_paragraph(doc, "网站地址：https://134139.xyz/")
    doc.add_page_break()


def add_role_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "角色"
    hdr[1].text = "主要职责"
    hdr[2].text = "主要入口"
    for cell in hdr:
        set_cell_shading(cell, "EEF6F1")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=11, bold=True, color="0F3D2B")

    rows = [
        ("普通用户", "注册登录、提交电子产品、查看审核进度、申请领取公示物品、积分兑换", "网站首页及“登录 / 注册”入口"),
        ("管理员", "审核物品、处理上门回收、审核申领、处理兑换、维护回收点、发布公告", "隐藏后台登录入口，如 index.php?page=管理员隐藏路径"),
    ]
    for role, duty, entry in rows:
        row = table.add_row().cells
        row[0].text = role
        row[1].text = duty
        row[2].text = entry
        for cell in row:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.35
                for run in paragraph.runs:
                    set_run_font(run, size=11)


def build_document():
    doc = Document()
    apply_base_styles(doc)
    add_cover(doc)

    add_heading(doc, "1. 平台概述", 1)
    add_body_paragraph(doc, "绿循校园是一个面向校园电子废物处理场景的网站平台，围绕“捐赠流转、固定回收点投放、上门回收”三类典型业务，构建了从用户提交、管理员审核、进度通知到积分激励的完整闭环。平台既服务于普通学生用户的便捷使用，也兼顾管理员日常审核、回收点维护和公告发布等管理需求。")
    add_body_paragraph(doc, "与传统的线下登记方式相比，该平台将物品提交、回收方式选择、审核反馈、申领流转和积分兑换整合到统一界面中，既降低了用户操作门槛，也提升了校园电子废物规范处理的组织效率。")
    add_shot_box(doc, "首页整体界面。", "网站顶部品牌区、导航栏、首页主标题区域、右侧平台功能简介区。", "本节下方，作为平台整体展示图。")

    add_heading(doc, "2. 角色划分与使用入口", 1)
    add_body_paragraph(doc, "平台包含两类核心使用角色，即普通用户与管理员。两类角色共用同一套网站基础结构，但在登录入口、可见页面、可执行操作和数据权限范围上存在明显差异。")
    add_role_table(doc)

    add_heading(doc, "2.1 普通用户入口", 2)
    add_body_paragraph(doc, "普通用户通过网站首页进入“登录 / 注册”页面，完成账号创建或账号登录后，即可访问提交物品、公示大厅、积分兑换和个人中心等前台模块。网站访问地址为 https://134139.xyz/。")

    add_heading(doc, "2.2 管理员入口", 2)
    add_body_paragraph(doc, "管理员后台采用隐藏访问路径，不直接在网站首页公开展示。管理员需要通过特定链接进入登录页，输入管理员手机号和密码后方可进入后台。当前管理员登录地址为 https://134139.xyz/index.php?page=xmu-greenloop-admin-6f9c2d71。该设计能够降低后台入口被随意访问的风险，并增强系统安全性。")
    add_tip_box(doc, "如果你后续用于答辩或提交材料，可在此处补充“管理员默认入口示例”与“管理员默认账号信息”，但正式公开版本建议隐藏敏感路径与密码。")
    add_shot_box(doc, "登录 / 注册页，以及管理员登录页。", "用户登录标签页、注册标签页、管理员隐藏登录入口页面。", "本节下方，展示双角色入口差异。")

    add_heading(doc, "3. 共用功能机制", 1)
    add_body_paragraph(doc, "为避免文档重复，本节先对用户端与管理员端均会接触的共用功能机制进行统一说明，包括身份认证、信息填写与图片上传、状态流转、通知反馈和积分规则等内容。")

    add_heading(doc, "3.1 身份认证与权限控制", 2)
    add_body_paragraph(doc, "平台通过手机号与密码实现身份认证。普通用户可自助注册账号并完成登录，管理员则通过隐藏后台入口进行登录。系统按照角色控制页面访问权限：普通用户只能查看与本人相关的提交、申请、消息和兑换记录，管理员则能够进入后台执行审核、维护和发布操作。")

    add_heading(doc, "3.2 信息填写与图片上传", 2)
    add_body_paragraph(doc, "在提交电子产品时，平台要求用户填写物品名称、类别、品牌或型号、成色或状态、文字说明、实物图片以及处理方式。图片上传支持 JPG、PNG、WebP、GIF 等常见格式，系统对文件大小进行限制，以保证展示效果和上传效率。通过结构化信息填写，可以帮助管理员快速判断物品状态，也方便后续公示、回收和统计。")
    add_shot_box(doc, "“提交物品”页面上半部分。", "物品名称、类别、品牌型号、文字说明、图片上传、处理方式等基础字段。", "本小节下方，说明平台统一的信息填写机制。")

    add_heading(doc, "3.3 状态流转与进度反馈", 2)
    add_body_paragraph(doc, "平台围绕不同处理方式设计了清晰的状态流转。所有物品在提交后都会先进入“待审核”状态，经管理员处理后再分别进入公示、待投放、待上门回收、已匹配或已完成等状态。该设计可以让用户与管理员对每条记录所处阶段形成统一认知，减少线下沟通成本。")
    add_body_paragraph(doc, "其中，捐赠流程主要经历“待审核—公示中—已匹配—已完成”；固定回收点流程主要经历“待审核—待投放—已完成”；上门回收流程主要经历“待审核—待上门回收—已完成”。")

    add_heading(doc, "3.4 站内通知机制", 2)
    add_body_paragraph(doc, "平台通过站内通知同步关键处理结果，包括物品审核结果、是否进入公示或回收流程、申领审核结果、积分到账提醒以及兑换申请处理结果等。通知机制使用户无需反复询问管理员，也有助于管理员将处理意见标准化、留痕化。")

    add_heading(doc, "3.5 积分激励机制", 2)
    add_body_paragraph(doc, "平台将积分与真实处理结果绑定，用于鼓励用户参与校园电子废物规范处理。当前规则为：捐赠审核通过后奖励 5 积分，固定回收点审核通过后奖励 5 积分，上门回收在管理员确认完成后奖励 5 积分。用户提交积分兑换申请后，系统会先暂扣相应积分；若管理员驳回，则自动退回积分。")
    add_shot_box(doc, "个人中心通知区域、积分兑换页面中的积分规则区域。", "站内通知列表、当前积分显示、积分规则说明卡片。", "本节末尾，说明“状态 + 通知 + 积分”的闭环机制。")

    doc.add_page_break()

    add_heading(doc, "4. 用户端功能说明", 1)
    add_body_paragraph(doc, "用户端主要面向普通学生或校内教职工，核心目标是帮助其完成电子产品捐赠、回收投放或上门预约处理，并随时查询进度与处理结果。")

    add_heading(doc, "4.1 账号注册与登录", 2)
    add_body_paragraph(doc, "普通用户首次使用平台时，需要在“登录 / 注册”页面完成账号注册。注册时需填写手机号、短信验证码、密码、昵称和所属校区。完成注册后可自动进入已登录状态，后续使用手机号和密码即可再次登录。该设计兼顾了实名性、便捷性与校园场景的基本管理要求。")
    add_shot_box(doc, "用户注册界面、用户登录界面。", "手机号、验证码、密码、昵称、所属校区字段，以及登录表单。", "本小节下方。")

    add_heading(doc, "4.2 提交待处理电子产品", 2)
    add_body_paragraph(doc, "登录后，用户可进入“提交物品”页面提交电子产品信息。该页面是用户端最核心的业务入口，既支持可继续使用设备的捐赠提交，也支持损坏设备的固定回收点投放或上门回收预约。通过统一提交入口，平台避免了多入口造成的理解困难，使用户先填写基础信息，再根据处理方式补充对应字段。")

    add_heading(doc, "4.2.1 捐赠", 3)
    add_body_paragraph(doc, "当设备仍可继续使用时，用户可选择“捐赠”方式，并补充捐赠说明和希望优先面向的人群。管理员审核通过后，该物品将进入公示大厅，供其他有需要的用户发起申领。该功能适用于闲置但仍有使用价值的电子设备，有助于在校内实现资源再利用。")

    add_heading(doc, "4.2.2 投放固定回收点", 3)
    add_body_paragraph(doc, "当设备已经损坏或老旧，但用户可以自行携带时，可选择“投放固定回收点”方式。此时页面会要求用户补充校区、园区、具体点位和预计投放时间。管理员审核通过后，记录进入待投放状态，并纳入后续集中回收流程。该模式适合体积较小、便于携带的电子废弃物处理。")

    add_heading(doc, "4.2.3 上门回收", 3)
    add_body_paragraph(doc, "当设备不便自行搬运，或用户更希望管理员到宿舍上门处理时，可选择“上门回收”。页面会要求补充宿舍楼栋、楼层、房间号和可预约时间段。其中楼层设置为 1 至 13 层，房间号建议使用“1105/1209”格式，预约时间段采用平台提供的固定选项。管理员审核通过后，记录进入待上门回收状态；待上门回收完成后，系统再发放积分。")
    add_shot_box(doc, "提交物品页面的三种处理方式区域。", "1. 选择“捐赠”后的补充字段；2. 选择“投放固定回收点”后的校区、园区、点位字段；3. 选择“上门回收”后的楼栋、楼层、房间号、预约时段字段。", "本小节末尾，建议按三张图连续排版。")

    add_heading(doc, "4.3 公示大厅浏览与申领", 2)
    add_body_paragraph(doc, "对于已审核通过的捐赠物品，平台会展示在公示大厅中。用户可浏览物品名称、类别、状态描述、图片和基本说明，筛选符合自身需求的设备。若用户对某一物品有使用需求，可进入详情页填写申请原因或用途说明，提交申领申请并等待管理员审核。固定回收点和上门回收记录不进入公示大厅，因此不会开放申领。")
    add_shot_box(doc, "公示大厅列表页、物品详情页、申领表单区域。", "已公示物品卡片、详情介绍、申请用途填写框。", "本小节下方。")

    add_heading(doc, "4.4 个人中心", 2)
    add_body_paragraph(doc, "个人中心是用户查看本人业务进度的集中入口。用户可以在此查看我的提交、我的申请、站内通知、兑换记录和当前积分。通过该页面，用户不仅可以快速获知某条记录是否已审核、是否进入公示、是否待回收或已完成，还能集中查看管理员备注与后续安排。")
    add_shot_box(doc, "个人中心页面。", "我的提交列表、站内通知列表、当前积分显示、兑换记录区域。", "本小节下方。")

    add_heading(doc, "4.5 积分兑换", 2)
    add_body_paragraph(doc, "在积分兑换页面，用户可查看当前积分、可兑换商品、商品所需积分与库存情况。若积分满足要求，用户可提交兑换申请，系统会先暂扣积分并等待管理员确认。该功能使平台不仅具备回收管理属性，也具备一定的参与激励和活动运营能力。")
    add_shot_box(doc, "积分兑换页面。", "积分商品列表、当前积分显示、立即兑换按钮、积分规则区域。", "本小节下方。")

    add_heading(doc, "5. 管理员端功能说明", 1)
    add_body_paragraph(doc, "管理员端是平台业务闭环落地的核心支撑部分，负责将用户提交的信息转化为真实的审核、回收、匹配和发放动作。相较于用户端，管理员端更强调流程控制、状态维护和服务协调。")

    add_heading(doc, "5.1 后台概览", 2)
    add_body_paragraph(doc, "管理员登录后台后，可在概览页快速查看关键待办数据，包括待审核物品数量、待上门回收订单数量、待审核申请数量、待处理兑换数量以及当前启用中的回收点数量。该页面相当于后台的工作台，用于帮助管理员掌握当前系统压力和优先处理顺序。")
    add_shot_box(doc, "管理员后台概览页。", "顶部导航、概览统计卡片、最近提交区域。", "本小节下方。")

    add_heading(doc, "5.2 物品审核", 2)
    add_body_paragraph(doc, "物品审核模块用于处理所有新提交记录。管理员可查看物品名称、类别、状态、文字说明、图片和详情页，并在备注框中填写审核意见后选择“通过”或“驳回”。不同处理方式在审核通过后进入不同流程：捐赠进入公示大厅，固定回收点进入待投放状态，上门回收进入待上门回收状态。通过备注机制，管理员可以向用户传递进一步说明或修改建议。")
    add_shot_box(doc, "物品审核页面。", "待审核记录列表、审核备注框、通过 / 驳回按钮。", "本小节下方。")

    add_heading(doc, "5.3 上门回收订单处理", 2)
    add_body_paragraph(doc, "对于选择上门回收的记录，管理员可在专门的“上门回收”模块中查看待处理订单与完成记录。待处理订单会展示提交人昵称、脱敏手机号、宿舍信息和预约时段，便于管理员提前安排回收。完成线下回收后，管理员可填写回收备注并标记“已上门回收”，系统随即将状态更新为已完成并向用户发放积分。")
    add_shot_box(doc, "上门回收管理页面。", "待上门回收订单表格、宿舍信息、预约时段、完成操作按钮。", "本小节下方。")

    add_heading(doc, "5.4 申领申请审核", 2)
    add_body_paragraph(doc, "当公示大厅中的捐赠物品收到申领申请后，管理员可在“申请审核”模块统一处理。管理员需要查看申请物品、申请人信息以及用途说明，再决定是否通过。若通过，物品状态转为已匹配，并由系统向申请人发送领取说明，同时通知原提交者其物品已完成匹配。该模块承担着捐赠流转中的匹配决策作用。")
    add_shot_box(doc, "申请审核页面。", "申请列表、用途说明、管理员回复框、审核按钮。", "本小节下方。")

    add_heading(doc, "5.5 积分兑换处理", 2)
    add_body_paragraph(doc, "用户提交积分兑换申请后，管理员可在“积分兑换”模块中查看待处理记录。管理员可选择通过兑换并填写领取说明，也可驳回兑换申请。系统会在通过时将兑换状态调整为已发放，在驳回时自动退回积分并回补库存，从而保证兑换记录和商品库存的一致性。")
    add_shot_box(doc, "积分兑换处理页面。", "待处理兑换列表、商品名称、所需积分、领取说明填写区域、通过 / 驳回按钮。", "本小节下方。")

    add_heading(doc, "5.6 回收点管理", 2)
    add_body_paragraph(doc, "固定回收点信息的可用性直接影响用户能否顺利完成投放，因此平台提供了回收点管理模块。管理员可新增并维护回收点名称、所属校区、具体位置、开放时间和文字说明等内容。当前台用户选择“投放固定回收点”方式时，系统会依据这些配置项展示对应点位信息。")
    add_shot_box(doc, "回收点管理页面。", "新增回收点表单、回收点列表、开放时间和说明字段。", "本小节下方。")

    add_heading(doc, "5.7 公告管理", 2)
    add_body_paragraph(doc, "公告管理模块用于在首页同步展示重要通知信息，例如回收安排、商品领取说明、节假日服务调整或校园活动通知等。管理员可以新增公告并控制其显示状态，使首页既具备功能入口作用，也承担运营通知和活动发布作用。")
    add_shot_box(doc, "公告管理页面，以及首页公告栏。", "公告发布表单、公告列表、首页公告栏展示效果。", "本小节下方，可采用前后台对照排版。")

    add_heading(doc, "6. 业务闭环与系统特色", 1)
    add_body_paragraph(doc, "从整体功能来看，平台并非单一的信息展示网站，而是一个完整的校园电子废物治理业务系统。其核心价值体现在以下几个方面：")
    add_bullets(doc, [
        "将捐赠流转、固定回收点投放和上门回收三种场景统一整合到同一平台。",
        "通过管理员审核与流程分发，保证信息真实、流程清晰、责任明确。",
        "通过公示大厅与申领审核，实现校内可用电子设备的再利用。",
        "通过站内通知与个人中心，实现全过程留痕与用户可感知反馈。",
        "通过积分兑换机制，增强用户参与校园环保行动的积极性。",
    ])
    add_shot_box(doc, "选择 3 至 4 张最能代表闭环的界面拼图。", "首页、提交物品、公示大厅、个人中心或后台审核页。", "本节下方，作为“系统整体闭环示意图”。")

    add_heading(doc, "7. 建议配图清单汇总", 1)
    add_body_paragraph(doc, "为方便后续补图，本文档建议至少准备以下界面截图：")
    add_numbered(doc, [
        "首页整体界面。",
        "用户登录 / 注册页面。",
        "管理员隐藏登录页面。",
        "提交物品页面基础信息区域。",
        "捐赠表单补充字段界面。",
        "固定回收点表单补充字段界面。",
        "上门回收表单补充字段界面。",
        "公示大厅列表页。",
        "物品详情页与申领表单。",
        "个人中心页面。",
        "积分兑换页面。",
        "后台概览页。",
        "物品审核页面。",
        "上门回收管理页面。",
        "申请审核页面。",
        "积分兑换处理页面。",
        "回收点管理页面。",
        "公告管理页面与首页公告栏。",
    ])
    add_body_paragraph(doc, "建议优先保留“首页—提交物品—公示大厅—个人中心—后台审核”这一组主流程截图，可更直观地体现系统闭环。")

    add_heading(doc, "8. 结语", 1)
    add_body_paragraph(doc, "本功能文档对绿循校园网站的用户端与管理员端能力进行了统一梳理，重点强调了共用机制、角色分工、主要业务流程与后续配图位置。后续你只需要按照文中“建议截图”提示逐项截取界面，并插入对应段落下方，即可快速整理成完整的图文版项目说明文档或答辩展示材料。")

    try:
        doc.save(str(OUTPUT_PATH))
    except PermissionError:
        fallback_path = OUTPUT_PATH.with_name(OUTPUT_PATH.stem + "-更新版" + OUTPUT_PATH.suffix)
        doc.save(str(fallback_path))


if __name__ == "__main__":
    build_document()
