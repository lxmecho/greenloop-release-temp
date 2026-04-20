from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOC_PATH = Path(r"d:\桌面\课程\节能减排大赛\网站功能文档-用户端与管理员端.docx")


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def main():
    doc = Document(str(DOC_PATH))
    inserted_site_line = False

    remove_prefixes = (
        "建议截图：",
        "文档写作提示：",
        "此外，概览页还提供演示数据初始化功能",
    )

    in_shot_list_section = False

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        if text == "7. 建议配图清单汇总":
            in_shot_list_section = True
            remove_paragraph(paragraph)
            continue

        if in_shot_list_section:
            if text == "8. 结语":
                paragraph.text = "7. 结语"
                in_shot_list_section = False
            else:
                remove_paragraph(paragraph)
            continue

        if text == "用户端与管理员端合并说明版（含配图预留提示）":
            paragraph.text = "用户端与管理员端合并说明版"
            continue

        if text == "用户端与管理员端合并说明版":
            continue

        if text.startswith("文档用途：") or text.startswith("文档说明："):
            remove_paragraph(paragraph)
            continue

        if text.startswith("项目名称：绿循校园"):
            next_paragraph = paragraph._element.getnext()
            next_text = ""
            if next_paragraph is not None:
                try:
                    next_text = "".join(next_paragraph.itertext()).strip()
                except Exception:
                    next_text = ""
            if not next_text.startswith("网站地址："):
                insert_paragraph_after(paragraph, "网站地址：https://134139.xyz/")
            inserted_site_line = True
            continue

        if text.startswith("网站地址："):
            paragraph.text = "网站地址：https://134139.xyz/"
            continue

        if text == "普通用户通过网站首页进入“登录 / 注册”页面，完成账号创建或账号登录后，即可访问提交物品、公示大厅、积分兑换和个人中心等前台模块。":
            paragraph.text = "普通用户通过网站首页进入“登录 / 注册”页面，完成账号创建或账号登录后，即可访问提交物品、公示大厅、积分兑换和个人中心等前台模块。网站访问地址为 https://134139.xyz/。"
            continue

        if text == "管理员后台采用隐藏访问路径，不直接在网站首页公开展示。管理员需要通过特定链接进入登录页，输入管理员手机号和密码后方可进入后台。该设计能够降低后台入口被随意访问的风险，并增强系统安全性。":
            paragraph.text = "管理员后台采用隐藏访问路径，不直接在网站首页公开展示。管理员需要通过特定链接进入登录页，输入管理员手机号和密码后方可进入后台。当前管理员登录地址为 https://134139.xyz/index.php?page=xmu-greenloop-admin-6f9c2d71。该设计能够降低后台入口被随意访问的风险，并增强系统安全性。"
            continue

        if text.startswith("本功能文档对绿循校园网站的用户端与管理员端能力进行了统一梳理"):
            paragraph.text = "本功能文档对绿循校园网站的用户端与管理员端能力进行了统一梳理，重点说明了共用机制、角色分工与主要业务流程。"
            continue

        if any(text.startswith(prefix) for prefix in remove_prefixes):
            remove_paragraph(paragraph)

    try:
        doc.save(str(DOC_PATH))
    except PermissionError:
        fallback = DOC_PATH.with_name(DOC_PATH.stem + "-最终版" + DOC_PATH.suffix)
        doc.save(str(fallback))


if __name__ == "__main__":
    main()
