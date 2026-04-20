from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"d:\桌面\课程\节能减排大赛")
OUTPUT_DIR = ROOT / "最终提交版"


def set_run_font(run, size=12, bold=False, color="000000"):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_border(paragraph, color="9DB8A7", size="8", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    for edge in ("top", "left", "bottom", "right"):
        element = p_bdr.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            p_bdr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), space)
        element.set(qn("w:color"), color)


def apply_base_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(12)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=20, bold=True, color="0F3D2B")

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(18)
        run2 = p2.add_run(subtitle)
        set_run_font(run2, size=11, color="5A6E63")


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=13.5, bold=True, color="0F3D2B")


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=12)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_run_font(run, size=12)


def add_shot_note(doc, pages, content):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_border(p, color="B7C9BC")

    r1 = p.add_run("建议截图页面：")
    set_run_font(r1, size=11, bold=True, color="0F3D2B")
    r2 = p.add_run(pages + "\n")
    set_run_font(r2, size=11)

    r3 = p.add_run("建议截取内容：")
    set_run_font(r3, size=11, bold=True, color="0F3D2B")
    r4 = p.add_run(content)
    set_run_font(r4, size=11)


def build_site_intro():
    doc = Document()
    apply_base_styles(doc)
    add_title(doc, "绿循校园网站简介", "提交材料精简版")

    add_heading(doc, "一、网站定位")
    add_paragraph(
        doc,
        "绿循校园是一个面向校园场景的电子废物规范处理网站，围绕“固定回收点投放、上门回收”两种处理方式，构建了从用户提交、管理员审核到结果反馈的基本业务闭环。",
    )

    add_heading(doc, "二、服务对象")
    add_paragraph(
        doc,
        "网站主要服务于校内学生和管理人员。普通用户可以在线提交损坏或老旧电子产品、查看处理进度；管理员可以在后台完成审核、回收安排、公告发布和积分管理等工作。",
    )

    add_heading(doc, "三、网站特点")
    add_bullets(
        doc,
        [
            "围绕校园电子废物处理场景设计，主题明确，使用门槛低。",
            "将回收点投放、上门回收、通知和积分激励整合到同一平台中。",
            "同时支持用户端操作与管理员端管理，具备完整的网站运行逻辑。",
        ],
    )
    add_shot_note(
        doc,
        "首页",
        "网站名称、主标题、两种处理方式入口或简介区域，以及页面整体视觉效果。",
    )

    add_heading(doc, "四、建设意义")
    add_paragraph(
        doc,
        "该网站能够为校园内老旧或损坏电子设备提供更规范的回收路径，提升电子废物处理效率，体现绿色校园建设的实际应用价值。",
    )

    return doc


def build_feature_intro():
    doc = Document()
    apply_base_styles(doc)
    add_title(doc, "绿循校园核心功能简介", "提交材料精简版")

    add_heading(doc, "一、链路起点：用户登录并提交回收信息")
    add_paragraph(
        doc,
        "用户通过手机号完成注册与登录后，进入“提交物品”页面填写电子产品名称、类别、状态、图片等基础信息，并选择“固定回收点投放”或“上门回收”中的一种处理方式。这一步是整条回收链路的起点。",
    )
    add_shot_note(
        doc,
        "登录 / 注册页，配合提交物品页",
        "手机号登录注册入口，以及物品基础信息表单、图片上传区域和处理方式选择区域。",
    )

    add_heading(doc, "二、链路第二步：管理员审核并生成处理安排")
    add_paragraph(
        doc,
        "用户提交后，记录先进入待审核状态。管理员在后台核对图片和文字说明，决定通过或驳回；审核通过后，系统生成物品编号，并根据用户所选方式分流到固定回收点流程或上门回收流程。",
    )
    add_shot_note(
        doc,
        "管理员后台“物品审核”页",
        "待审核列表、审核按钮、管理员备注输入区域，以及通过后的流程去向说明。",
    )

    add_heading(doc, "三、链路第三步：进入具体回收执行环节")
    add_paragraph(
        doc,
        "若用户选择固定回收点，需填写校区、园区、点位和预计投放时间，审核通过后按编号完成投递；若用户选择上门回收，需填写宿舍楼栋、楼层、房间号和预约时段，管理员再按预约安排上门取走。",
    )
    add_shot_note(
        doc,
        "提交物品页中的固定回收点区域和上门回收区域，必要时补充物品详情页",
        "固定回收点的校区、园区、点位、预计投放时间，以及上门回收的宿舍信息和预约时段。",
    )

    add_heading(doc, "四、链路第四步：完成投递或取走确认并进入核验")
    add_paragraph(
        doc,
        "固定回收点记录在用户确认“已投递”后进入统一回收和入仓核验阶段；上门回收记录在管理员与用户双方完成“已取走”确认后进入入仓核验阶段。该环节保证每条回收记录都有明确的过程反馈。",
    )
    add_shot_note(
        doc,
        "个人中心页，配合管理员后台“上门回收”页",
        "用户侧的状态变化、确认按钮，以及管理员侧的待处理订单列表和处理按钮。",
    )

    add_heading(doc, "五、链路终点：站内通知与积分反馈")
    add_paragraph(
        doc,
        "管理员完成最终核验后，系统会通过站内消息向用户反馈处理结果，并按照参考价自动换算和发放积分。用户可在个人中心查看通知和积分变化，并在积分商城继续完成兑换，形成完整闭环。",
    )
    add_shot_note(
        doc,
        "个人中心页，或积分商城页",
        "站内消息列表、当前积分显示、积分到账结果，以及积分兑换相关区域。",
    )

    return doc


def build_feature_doc():
    doc = Document()
    apply_base_styles(doc)
    add_title(doc, "绿循校园网站功能文档", "用户端与管理员端精简说明版")

    add_heading(doc, "一、平台概述")
    add_paragraph(
        doc,
        "绿循校园是一个面向校园电子废物处理场景的网站平台，围绕固定回收点投放和上门回收两类核心业务，构建了从用户提交、管理员审核、进度通知到积分激励的完整闭环。",
    )
    add_shot_note(
        doc,
        "首页",
        "网站主标题、两种处理方式简介区、导航栏和公告区。",
    )

    add_heading(doc, "二、角色与入口")
    add_paragraph(
        doc,
        "平台包含普通用户与管理员两类角色。普通用户通过首页的“登录 / 注册”入口进入前台模块；管理员通过隐藏后台入口登录后，可处理审核、回收流程、公告和积分兑换。",
    )
    add_shot_note(
        doc,
        "登录 / 注册页，以及管理员登录页",
        "普通用户登录注册区域、管理员隐藏登录入口页面。",
    )

    add_heading(doc, "三、用户端核心功能")
    add_bullets(
        doc,
        [
            "提交物品：填写物品名称、类别、状态、图片和处理方式，是最核心的业务入口。",
            "固定回收点：补充校区、园区、点位和预计投放时间，审核通过后按编号完成投递。",
            "上门回收：补充宿舍信息和预约时段，管理员上门取走后进入核验流程。",
            "个人中心与积分兑换：查看提交记录、站内通知、当前积分和兑换记录。",
        ],
    )
    add_shot_note(
        doc,
        "提交物品页、个人中心页、积分商城页",
        "基础信息表单、两种处理方式区域、消息列表、当前积分和兑换商品区域。",
    )

    add_heading(doc, "四、管理员端核心功能")
    add_bullets(
        doc,
        [
            "物品审核：核对用户提交的图片和说明，决定通过或驳回。",
            "上门回收处理：查看待上门回收订单，完成取走确认并推进核验。",
            "回收点与公告管理：维护回收点信息和首页公告内容。",
            "积分兑换处理：审核兑换申请，决定发放或驳回。",
        ],
    )
    add_shot_note(
        doc,
        "管理员后台“物品审核”页，必要时补充“上门回收”页",
        "待审核列表、处理按钮、回收流程状态、公告或兑换管理区域。",
    )

    add_heading(doc, "五、总结")
    add_paragraph(
        doc,
        "系统已形成“提交 - 审核 - 回收 - 通知 - 积分兑换”的业务闭环，能够较直观展示校园电子废物处理场景下的网站设计思路和实现效果。",
    )

    return doc


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    root_feature_doc_path = ROOT / "网站功能文档-用户端与管理员端.docx"
    feature_doc_path = OUTPUT_DIR / "01-网站功能文档.docx"
    site_intro_path = OUTPUT_DIR / "06-网站简介.docx"
    feature_intro_path = OUTPUT_DIR / "07-核心功能简介.docx"

    feature_doc = build_feature_doc()
    feature_doc.save(root_feature_doc_path)
    feature_doc.save(feature_doc_path)
    build_site_intro().save(site_intro_path)
    build_feature_intro().save(feature_intro_path)

    print(root_feature_doc_path)
    print(feature_doc_path)
    print(site_intro_path)
    print(feature_intro_path)


if __name__ == "__main__":
    main()
